from pathlib import Path

import fitz
import pytest

from app import ocr
from app.extractors import extract_docx_pages, extract_pdf_pages, extract_txt_pages


def _make_text_pdf(path: Path, text: str) -> None:
    doc = fitz.open()
    doc.new_page().insert_text((72, 72), text, fontsize=14)
    doc.save(path)
    doc.close()


def _make_scanned_pdf(path: Path, text: str) -> None:
    """A PDF page containing only a rendered image of `text` -- no text
    layer at all, simulating a scanned document (see app/ocr.py)."""
    src = fitz.open()
    src_page = src.new_page()
    src_page.insert_text((72, 200), text, fontsize=48)
    pix = src_page.get_pixmap(matrix=fitz.Matrix(3, 3))
    src.close()

    out = fitz.open()
    out.new_page().insert_image(out[0].rect, pixmap=pix)
    out.save(path)
    out.close()


def test_extract_txt_pages_returns_single_page(tmp_path):
    path = tmp_path / "notes.txt"
    path.write_text("Handrails shall be continuous.", encoding="utf-8")

    pages = extract_txt_pages(path)

    assert pages == ["Handrails shall be continuous."]


def test_extract_docx_pages_joins_paragraphs_as_single_page(tmp_path):
    import docx

    path = tmp_path / "addendum.docx"
    document = docx.Document()
    document.add_paragraph("Fire Door Requirements")
    document.add_paragraph("Corridor doors shall have a twenty minute rating.")
    document.save(path)

    pages = extract_docx_pages(path)

    assert len(pages) == 1
    assert "Fire Door Requirements" in pages[0]
    assert "Corridor doors shall have a twenty minute rating." in pages[0]


def test_extract_pdf_pages_with_text_layer_skips_ocr(tmp_path):
    path = tmp_path / "text.pdf"
    _make_text_pdf(path, "Fire doors shall be rated for twenty minutes.")

    pages, ocr_status = extract_pdf_pages(path)

    assert ocr_status == "not_needed"
    assert "Fire doors shall be rated for twenty minutes." in pages[0]


def test_extract_pdf_pages_ocrs_scanned_page(tmp_path):
    if not ocr.is_available():
        pytest.skip("Tesseract not installed on this machine")

    path = tmp_path / "scanned.pdf"
    _make_scanned_pdf(path, "SAMPLE TEXT")

    pages, ocr_status = extract_pdf_pages(path)

    assert ocr_status == "used"
    assert "SAMPLE" in pages[0].upper()


def test_extract_pdf_pages_reports_unavailable_when_tesseract_missing(tmp_path, monkeypatch):
    monkeypatch.setattr(ocr, "is_available", lambda: False)

    path = tmp_path / "scanned.pdf"
    _make_scanned_pdf(path, "SAMPLE TEXT")

    pages, ocr_status = extract_pdf_pages(path)

    assert ocr_status == "unavailable"
    assert pages[0] == ""
